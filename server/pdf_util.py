import datetime
import random
from pdf2docx import Converter
from docx2pdf import convert
from docx import Document
import re
import hashlib
import os
import shutil
from docxcompose.composer import Composer

def convert_pdf_to_docx(pdf_path, docx_path):
    if not os.path.isfile(docx_path):
        print("Converting " + pdf_path + " to " + docx_path)
        pdf = Converter(pdf_path)
        pdf.convert(docx_path)
        pdf.close()
        print("Converted " + pdf_path + " to " + docx_path)
    else:
        print("File already exists: " + docx_path)

def convert_docx_to_pdf(docx_path, pdf_path):
    if not os.path.isfile(pdf_path):
        print("Converting " + docx_path + " to " + pdf_path)
        convert(docx_path, pdf_path)
        print("Converted " + docx_path + " to " + pdf_path)
    else:
        print("File already exists: " + pdf_path)

def replace_in_table(tables, search, replace):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                # Is there a table in this cell?
                if cell.tables:
                    replace_in_table(cell.tables, search, replace)
                # Is there a paragraph in this cell?
                if cell.paragraphs:
                    for paragraph in cell.paragraphs:
                        # Replace the text
                        paragraph.text = re.sub(search, replace, paragraph.text)

def replace_in_whole_document(document, search, replace):
    # Replace in the document
    for paragraph in document.paragraphs:
        if paragraph.text.__contains__(search):
            print("Replacing " + search + " with " + replace + " in " + paragraph.text)
        paragraph.text = re.sub(search, replace, paragraph.text)
    # Replace in the tables
    replace_in_table(document.tables, search, replace)

def replace_docx_visable_metadata(docx_path,number_of_questions):
    # Metadata
    document = Document(docx_path)
    composer = Composer(document)
    document.core_properties.author = "PastPaperGen"
    document.core_properties.title = "Paper Title"
    document.core_properties.subject = "Paper Subject"
    document.core_properties.created = datetime.datetime.now()
    document.core_properties.modified = datetime.datetime.now()
    # Titled metadata (visible)
    replace_in_whole_document(document, "{{date_generated}}", datetime.datetime.now().strftime("%d/%m/%Y"))
    # A hash of the time for development purposes
    replace_in_whole_document(document, "{{mark_scheme_number}}", hashlib.sha256(str(datetime.datetime.now()).encode('utf-8')).hexdigest()) 
    replace_in_whole_document(document, "{{duration.hours}}", "3")
    replace_in_whole_document(document, "{{duration.minutes}}", "15")
    print("Updated metadata for " + docx_path)
    docx_path = docx_path.replace("_blank-metadata.docx", ".docx")
    composer.save(docx_path)
    return docx_path

def add_questions_to_paper(docx_path,number_of_questions):
    paper = Document(docx_path)
    composer = Composer(paper)
    questions = []
    # Create a list of random 1's or 3's
    for i in range(int(number_of_questions)):
        # Get a random question from any subdirectory of the questions directory
        question_libary_list = os.listdir("server/static/docx/Maths/Libary/Questions")
        question_libary_list.remove("Admin")
        question_libary = random.choice(question_libary_list)
        question = random.choice(os.listdir("server/static/docx/Maths/Libary/Questions/" + question_libary))
        questions.append(question)
        question = Document("server/static/docx/Maths/Libary/Questions/" + question_libary + "/" + question)
        composer.append(question)
        print("Added question " + str(i + 1) + " to paper of ID " + docx_path)
    # Now go through the whole paper and replace the string "{{question_number}}" with the question number
    question_number = 1
    # Read the whole document
    for paragraph in paper.paragraphs:
        # Replace text
        if '{{question_number}}' in paragraph.text:
            inline = paragraph.runs
            # Loop added to work with runs (strings with same style)
            replaced = False
            for i in range(len(inline)):
                if '{{question_number}}' in inline[i].text:
                    text = inline[i].text.replace('{{question_number}}', str(question_number))
                    inline[i].text = text
                    question_number += 1
                    print("Replaced {{question_number}} with " + str(question_number - 1))
                    replaced = True
            if not replaced:
                print("Could not replace {{question_number}} for question " + str(question_number))
                question_number += 1
    # Save the document
    docx_path = docx_path.replace("Generating", "Generated")
    composer.save(docx_path)
    print("Saved paper to " + docx_path)
    return questions

def create_docx():
    dir = "server/static/docx/Maths/Generating/Papers/"
    template = Document("server/static/docx/Maths/Libary/Questions/Admin/OCR_cover.docx")
    composer = Composer(template)
    name = hashlib.sha256(str(datetime.datetime.now()).encode('utf-8')).hexdigest() + "_blank-metadata.docx"
    composer.save(dir + name)
    print("Created paper with ID " + name + " in " + dir)
    return dir + name

def generate_mark_scheme(docx_path, questions):
    # Create the mark scheme
    default_page = Document("server/static/docx/Maths/Libary/Questions/Admin/OCR_mark_scheme_default.docx")
    composer = Composer(default_page)
    # Add the questions
    for question in questions:
        paper_question = Document(question)
        composer.append(paper_question)
    # Replace the question numbers
    question_number = 1
    for paragraph in default_page.paragraphs:
        if '{{question_number}}' in paragraph.text:
            inline = paragraph.runs
            # Loop added to work with runs (strings with same style)
            replaced = False
            for i in range(len(inline)):
                if '{{question_number}}' in inline[i].text:
                    text = inline[i].text.replace('{{question_number}}', str(question_number))
                    inline[i].text = text
                    question_number += 1
                    print("Replaced {{question_number}} with " + str(question_number - 1))
                    replaced = True
            if not replaced:
                print("Could not replace {{question_number}} for question " + str(question_number))
                question_number += 1
    # Save the document
    docx_path = docx_path.replace("Generating", "Generated")

def delete_tmp_files(docx_path):
    os.remove(docx_path)
    print("Deleted " + docx_path)
    docx_path = docx_path.replace(".docx", "_blank-metadata.docx")
    os.remove(docx_path)
    print("Deleted " + docx_path)

def generate_paper(number_of_questions):
    print("Number of questions: " + str(number_of_questions))
    convert_pdf_to_docx("server/static/pdf/Maths/Libary/Questions/Admin/OCR_cover.pdf", "server/static/docx/Maths/Libary/Questions/Admin/OCR_cover.docx")
    filename = create_docx()
    generating_filename = replace_docx_visable_metadata(filename, str(number_of_questions))
    print("Current File Name: " + generating_filename)
    questions = add_questions_to_paper(generating_filename, number_of_questions)
    delete_tmp_files(generating_filename)
    # convert_docx_to_pdf(filename, filename.replace("docx", "pdf").replace("docx", "pdf"))
    generate_mark_scheme(filename, questions)
    return None

if __name__ == "__main__":
    generate_paper(50)